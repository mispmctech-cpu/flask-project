from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Set document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title Page
title = doc.add_heading('ZEROTH REVIEW', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.runs[0]
title_run.font.color.rgb = RGBColor(88, 28, 135)

subtitle = doc.add_heading('FACULTY PERFORMANCE MANAGEMENT SYSTEM', 1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Add submission details
details = doc.add_paragraph()
details.alignment = WD_ALIGN_PARAGRAPH.CENTER
details.add_run('Submitted by: ').bold = True
details.add_run('[Your Name/Team Name]\n')
details.add_run('Guide: ').bold = True
details.add_run('[Guide Name]\n')
details.add_run('Department: ').bold = True
details.add_run('[Your Department]\n')
details.add_run('Institution: ').bold = True
details.add_run('[Your College Name]')

doc.add_page_break()

# 1. TITLE
doc.add_heading('1. TITLE', 1)
p = doc.add_paragraph()
p.add_run('Faculty Performance Management & IQAC Compliance Tracking System').bold = True

doc.add_page_break()

# 2. ABSTRACT
doc.add_heading('2. ABSTRACT', 1)
doc.add_paragraph(
    'A web-based application designed to automate and streamline faculty performance monitoring in alignment with '
    'IQAC/NAAC requirements. The system provides role-based access for Faculty, HODs, IQAC, and Principal to submit, '
    'verify, and monitor various activities including core scope (15 monthly activities), yearly mandatory portfolios '
    '(AP/ASP/Prof specific), department portfolios (8 forms), and institution portfolios (17 forms). Built using Flask '
    'and Supabase, it features real-time dashboards, automated calculations, verification workflows, and comprehensive '
    'reporting capabilities to eliminate manual paperwork and ensure compliance.'
)

doc.add_page_break()

# 3. OBJECTIVES
doc.add_heading('3. OBJECTIVES', 1)
objectives = [
    'Automate faculty activity tracking across teaching, research, and administrative domains',
    'Ensure IQAC/NAAC compliance through systematic monitoring',
    'Implement multi-level verification workflow (Faculty → HOD → IQAC → Principal)',
    'Provide real-time dashboards for performance monitoring',
    'Eliminate manual paperwork and enable data-driven decision making',
    'Maintain complete audit trails for transparency and accountability'
]
for i, obj in enumerate(objectives, 1):
    doc.add_paragraph(f'{i}. {obj}', style='List Number')

doc.add_page_break()

# 4. MODULES
doc.add_heading('4. MODULES', 1)

modules = [
    ('Authentication Module', 'Multi-role login system (Faculty, HOD, IQAC, Principal, Admin)'),
    ('Faculty Core Scope Module', '15 monthly activities tracking with percentage calculation'),
    ('Yearly Mandatory Scope Module', 'Designation-specific forms (AP: 8 activities, ASP: 9, Prof: 9)'),
    ('Department Portfolio Module', '8 monthly forms for department-level responsibilities'),
    ('Institution Portfolio Module', '17 forms with varied frequencies (weekly/monthly/yearly)'),
    ('HOD Verification Module', 'Review, approve/reject submissions with remarks'),
    ('IQAC Dashboard Module', 'Monitor all activities with filters and Excel export'),
    ('Principal Dashboard Module', 'Institution-wide overview with analytics'),
    ('Admin Module', 'User management, form control, database diagnostics'),
    ('Notification Module', 'Real-time alerts for submissions and verifications'),
    ('Reports & Analytics Module', 'Performance reports with PDF/Excel export'),
    ('Document Management Module', 'Secure file upload and storage system')
]

for i, (module, desc) in enumerate(modules, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(f'{module}: ').bold = True
    p.add_run(desc)

doc.add_page_break()

# 5. FLOW DIAGRAM
doc.add_heading('5. SYSTEM FLOW DIAGRAM', 1)
flow = [
    'User Login → Role-based Access',
    'Faculty → Fill & Submit Forms (Core/Yearly/Department/Institution)',
    'Database Storage → Supabase stores submissions',
    'HOD → Review & Verify/Reject with remarks',
    'Notification → Alerts for status updates',
    'IQAC/Principal → Monitor dashboards and analytics',
    'Reports → Generate PDF/Excel exports'
]
for step in flow:
    doc.add_paragraph(step, style='List Bullet')

doc.add_page_break()

# 6. ARCHITECTURE
doc.add_heading('6. ARCHITECTURE DIAGRAM', 1)

doc.add_heading('Three-Tier Architecture:', 2)
doc.add_paragraph('1. Presentation Layer: Web Browsers (HTML5, CSS3, JavaScript, Tailwind CSS)')
doc.add_paragraph('2. Application Layer: Flask Web Application (Python 3.x, routing, business logic)')
doc.add_paragraph('3. Database Layer: Supabase (PostgreSQL database + Storage buckets)')

doc.add_heading('Key Components:', 2)
components = [
    'Frontend: HTML5, CSS3, JavaScript, Tailwind CSS, Chart.js, SheetJS',
    'Backend: Python Flask, Werkzeug, Gunicorn',
    'Database: Supabase (PostgreSQL) with 30+ tables',
    'Storage: Supabase Storage for file uploads',
    'Security: Role-based access control + Row-Level Security (RLS)'
]
for comp in components:
    doc.add_paragraph(comp, style='List Bullet')

doc.add_page_break()

# 7. REFERENCES
doc.add_heading('7. REFERENCES', 1)
references = [
    'Flask Documentation - https://flask.palletsprojects.com/',
    'Supabase Documentation - https://supabase.com/docs',
    'PostgreSQL Documentation - https://www.postgresql.org/docs/',
    'Tailwind CSS - https://tailwindcss.com/docs',
    'NAAC Assessment Framework - https://www.naac.gov.in/',
    'IQAC Best Practices Manual - UGC Guidelines',
    'OWASP Web Security Standards - https://owasp.org/',
    '"Web Application Development with Flask" - Miguel Grinberg'
]
for ref in references:
    doc.add_paragraph(ref, style='List Bullet')

# Save the document
output_path = r'c:\Users\91636\Downloads\Flask-project-main\PROJECT_REVIEW_SLIDES.docx'
doc.save(output_path)
print(f'✅ Concise Word document created: {output_path}')
